"""Extract plain text from uploaded documents (PDF / DOCX / TXT-ish)."""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class UnsupportedDocument(Exception):
    """Raised when we cannot extract text from the given file."""


def extract_pages(path: str, content_type: str | None = None) -> list[tuple[int | None, str]]:
    """Return ``[(page_number, text), ...]``.

    PDFs yield one entry per page (1-based); other formats yield a single entry
    with ``page_number=None``.
    """
    p = Path(path)
    suffix = p.suffix.lower()

    if suffix == ".pdf" or (content_type or "").endswith("pdf"):
        return _extract_pdf(p)
    if suffix == ".docx" or "wordprocessingml" in (content_type or ""):
        return [(None, _extract_docx(p))]
    if suffix in {".txt", ".md", ".csv", ".json", ""} or (content_type or "").startswith("text/"):
        return [(None, p.read_text(encoding="utf-8", errors="replace"))]
    raise UnsupportedDocument(f"Unsupported file type: {p.name} ({content_type})")


def _extract_pdf(path: Path) -> list[tuple[int | None, str]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[tuple[int | None, str]] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            pages.append((i, page.extract_text() or ""))
        except Exception:  # noqa: BLE001 - a broken page shouldn't kill ingestion
            logger.warning("Failed to extract text from page %s of %s", i, path.name)
            pages.append((i, ""))
    return pages


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
